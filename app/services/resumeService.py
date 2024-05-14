from docx import Document
from app.models.resume import Resume

def createResume(data: Resume) -> str:
  document = Document()
  document.add_heading('Currículo', 0)

  document.add_heading('Nome', level=1)
  document.add_paragraph(data.name)

  document.add_heading('Contato', level=1)
  document.add_paragraph(data.contact)

  document.add_heading('Experiência Profissional', level=1)
  for job in data.experience:
    document.add_heading(job.title, level=2)
    document.add_paragraph(job.company)
    document.add_paragraph(job.description)

  document.add_heading('Formação Acadêmica', level=1)
  
  for edu in data.education:
    document.add_heading(edu.degree, level=2)
    document.add_paragraph(edu.institution)
    document.add_paragraph(edu.description)

  file_path = 'resume.docx'
  document.save(file_path)
  return file_path